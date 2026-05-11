#!/usr/bin/env python
"""Generate ARPM research docx files from UTF-8 source text."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_base_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.38


def add_center_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(15)


def add_meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(14)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text)


def build_formal_doc(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_base_style(doc)

    title = lines[0].strip() if lines else ""
    subtitle = lines[1].strip() if len(lines) > 1 else ""
    meta = lines[2].strip() if len(lines) > 2 else ""

    abstract_lines: list[str] = []
    keywords = ""
    body_start = 0
    in_abstract = False
    for idx, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("摘要："):
            in_abstract = True
            value = stripped.replace("摘要：", "", 1).strip()
            if value:
                abstract_lines.append(value)
        elif stripped.startswith("关键词："):
            in_abstract = False
            keywords = stripped.replace("关键词：", "", 1).strip()
        elif stripped.startswith("# "):
            body_start = idx
            break
        elif in_abstract and stripped:
            abstract_lines.append(stripped)

    abstract = "\n".join(abstract_lines).strip()

    add_center_title(doc, title, subtitle)
    if meta:
        add_meta(doc, meta)

    add_heading(doc, "摘要")
    add_body_paragraph(doc, abstract)

    if keywords:
        p = doc.add_paragraph()
        p.add_run("关键词：").bold = True
        p.add_run(keywords)

    current_heading = None
    for line in lines[body_start:]:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            current_heading = line[2:].strip()
            add_heading(doc, current_heading)
        else:
            add_body_paragraph(doc, line)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_brief_doc(source: Path, output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    doc = Document()
    set_base_style(doc)

    add_center_title(doc, "ARPM 研究成果与工程化简述")
    add_meta(doc, "整理时间：2026-04-14\n材料来源：research\\杂项 下的早期 tex、docx 与 notes 文档。")

    for block in text.split("\n\n"):
        block = block.strip()
        if not block or block.startswith("ARPM 研究成果与工程化简述") or set(block) == {"="}:
            continue
        add_body_paragraph(doc, block)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--brief-source", required=True)
    parser.add_argument("--brief-output", required=True)
    parser.add_argument("--formal-source", required=True)
    parser.add_argument("--formal-output", required=True)
    parser.add_argument("--brief-copy")
    parser.add_argument("--formal-copy")
    args = parser.parse_args()

    build_brief_doc(Path(args.brief_source), Path(args.brief_output))
    build_formal_doc(Path(args.formal_source), Path(args.formal_output))

    if args.brief_copy:
        Path(args.brief_copy).write_bytes(Path(args.brief_output).read_bytes())
    if args.formal_copy:
        Path(args.formal_copy).write_bytes(Path(args.formal_output).read_bytes())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
